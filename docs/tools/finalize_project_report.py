from __future__ import annotations

import argparse
import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FOREST = "173629"
PAPER = "F5F0E6"
SAND = "DCCDB7"
WHITE = "FFFFFF"
INK = "181A18"


def find_paragraph(document: DocumentObject, exact_text: str):
    for paragraph in document.paragraphs:
        if " ".join(paragraph.text.split()) == exact_text:
            return paragraph
    raise ValueError(f"Không tìm thấy đoạn neo: {exact_text}")


def clear_between(start_paragraph, end_paragraph) -> None:
    current = start_paragraph._p.getnext()
    while current is not None and current is not end_paragraph._p:
        following = current.getnext()
        current.getparent().remove(current)
        current = following


class Writer:
    def __init__(self, document: DocumentObject, anchor) -> None:
        self.document = document
        self.anchor = anchor

    def _move_paragraph(self, paragraph):
        self.anchor._p.addprevious(paragraph._p)
        return paragraph

    def paragraph(
        self,
        text: str = "",
        *,
        style: str = "Normal",
        bold_prefix: str | None = None,
        italic: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    ):
        paragraph = self.document.add_paragraph(style=style)
        paragraph.alignment = alignment
        if bold_prefix and text.startswith(bold_prefix):
            first = paragraph.add_run(bold_prefix)
            first.bold = True
            paragraph.add_run(text[len(bold_prefix) :])
        else:
            run = paragraph.add_run(text)
            run.italic = italic
        return self._move_paragraph(paragraph)

    def heading(self, text: str, level: int = 2):
        style = "CAP2" if level == 2 else "CAP3"
        paragraph = self.document.add_paragraph(text, style=style)
        paragraph.paragraph_format.keep_with_next = True
        return self._move_paragraph(paragraph)

    def bullet(self, label: str, body: str):
        paragraph = self.document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.first_line_indent = Cm(-0.5)
        paragraph.add_run("• ")
        label_run = paragraph.add_run(label)
        label_run.bold = True
        paragraph.add_run(body)
        return self._move_paragraph(paragraph)

    def page_break(self):
        paragraph = self.document.add_paragraph(style="Normal")
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return self._move_paragraph(paragraph)

    def caption(self, text: str, *, table: bool = False):
        style = "VN_TableCaption" if table else "VN_Caption"
        paragraph = self.document.add_paragraph(text, style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = table
        return self._move_paragraph(paragraph)

    def figure(self, image_path: Path, caption: str, width_cm: float = 15.2):
        paragraph = self.document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
        self._move_paragraph(paragraph)
        self.caption(caption)

    def table(
        self,
        caption: str,
        headers: list[str],
        rows: Iterable[Iterable[str]],
        *,
        widths_cm: list[float] | None = None,
        font_size: float = 10.5,
    ):
        self.caption(caption, table=True)
        row_values = [list(row) for row in rows]
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = True
        header = table.rows[0]
        for index, value in enumerate(headers):
            cell = header.cells[index]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), FOREST)
            cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
        header_props = header._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        header_props.append(repeat)

        for row_value in row_values:
            row = table.add_row()
            for index, value in enumerate(row_value):
                cell = row.cells[index]
                cell.text = str(value)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(font_size)
                        run.font.color.rgb = RGBColor.from_string(INK)
        if widths_cm:
            for row in table.rows:
                for index, width in enumerate(widths_cm):
                    row.cells[index].width = Cm(width)
        self.anchor._p.addprevious(table._tbl)
        return table

    def code(self, content: str):
        table = self.document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F4F1EA")
        cell._tc.get_or_add_tcPr().append(shading)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(content)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        self.anchor._p.addprevious(table._tbl)
        return table


def write_chapter_three(writer: Writer, diagrams: Path) -> None:
    writer.paragraph(
        "Chương này phân tích và thiết kế hệ thống DẤU VỊ – Vietnam Traceable Coffee "
        "Collection trên cơ sở mã nguồn đã xây dựng. Phạm vi hiện thực gồm storefront "
        "Next.js, REST API FastAPI, PostgreSQL, xác thực bằng session, đơn hàng COD demo "
        "và chatbot hybrid RAG. Những chức năng quản trị chưa có trong bản demo được ghi "
        "rõ là thiết kế mở rộng, tránh nhầm lẫn giữa kiến trúc mục tiêu và kết quả hiện tại."
    )

    writer.heading("3.1. Yêu cầu chức năng của hệ thống")
    writer.paragraph(
        "Các yêu cầu chức năng được nhóm theo hành trình khám phá sản phẩm, truy xuất, tư "
        "vấn, mua hàng và vận hành dữ liệu. Frontend không import trực tiếp mock data từ "
        "page; mọi màn hình đọc dữ liệu qua ProductRepository, TraceabilityRepository hoặc "
        "AdvisorRepository và lựa chọn implementation mock/HTTP bằng data-source factory."
    )
    writer.table(
        "Bảng 3.1. Yêu cầu chức năng và trạng thái hiện thực",
        ["Mã", "Yêu cầu", "Kết quả mong đợi", "Trạng thái"],
        [
            ("FR01", "Khám phá sản phẩm", "Xem đủ 6 SKU, tìm kiếm, lọc, sắp xếp và giữ bộ lọc trên URL.", "Đã triển khai"),
            ("FR02", "Chi tiết sản phẩm", "Chọn format, trọng lượng, kiểu xay, số lượng; xem passport và sản phẩm liên quan.", "Đã triển khai"),
            ("FR03", "Giỏ hàng", "Zustand persist bằng khóa vtc-cart-v1, sửa số lượng, xóa và tính tổng.", "Đã triển khai"),
            ("FR04", "Checkout COD", "Validate người nhận; backend tự tính giá/phí và tạo đơn demo có idempotency.", "Đã triển khai"),
            ("FR05", "Truy xuất", "Tra mã lô, hiển thị evidence/timeline và disclosure Demo Data bắt buộc.", "Đã triển khai"),
            ("FR06", "Coffee Advisor", "Quiz 6 bước, chấm điểm có trọng số và trả top 3 kèm lý do.", "Đã triển khai"),
            ("FR07", "Coffee Assistant", "Hội thoại RAG trên 6 sản phẩm; có grounding, fallback và action hợp lệ.", "Đã triển khai"),
            ("FR08", "Tài khoản", "Đăng ký, đăng nhập, đọc phiên và đăng xuất bằng cookie HttpOnly.", "Đã triển khai"),
            ("FR09", "Quản trị", "CRUD catalog/đơn/knowledge base, RBAC ADMIN và re-index.", "Thiết kế mở rộng"),
        ],
        widths_cm=[1.2, 3.1, 8.1, 2.4],
    )

    writer.heading("3.2. Yêu cầu phi chức năng của hệ thống")
    writer.paragraph(
        "Yêu cầu phi chức năng xác định tiêu chuẩn chất lượng và các điều kiện bất biến khi "
        "triển khai. Các ngưỡng hiệu năng dưới đây là mục tiêu thiết kế; kết quả kiểm thử "
        "thực tế được trình bày ở Chương 4."
    )
    writer.table(
        "Bảng 3.2. Yêu cầu phi chức năng",
        ["Nhóm", "Yêu cầu thiết kế", "Cơ chế đáp ứng"],
        [
            ("Hiệu năng", "Trang chính phản hồi trong khoảng 3 giây trên mạng thông thường; API đọc mục tiêu dưới 500 ms; chatbot có timeout 20 giây.", "RSC mặc định, tối ưu asset, GZip, index SQL/HNSW, top-k giới hạn và fallback khi Groq chậm/lỗi."),
            ("Bảo mật", "Không lưu mật khẩu/session token dạng rõ; không tin giá từ client; secret không xuất hiện ở frontend.", "Argon2id, SHA-256 token hash, cookie HttpOnly/Secure/SameSite, Origin check, rate limit, CORS allow-list."),
            ("Khả dụng", "Hoạt động ở 375, 768, 1280 và 1440 px; thao tác bàn phím và focus rõ.", "Responsive layout, hit area tối thiểu, dialog focus trap, reduced-motion và error/empty/loading states."),
            ("Độ tin cậy", "Không tạo đơn trùng; tổng tiền lấy từ server; dịch vụ chỉ ready khi DB/vector bắt buộc sẵn sàng.", "SQL transaction, Idempotency-Key, snapshot order item, health/live, health/ready và health/rag."),
            ("Mở rộng", "Tách frontend/backend và cho phép thay nguồn dữ liệu hoặc model mà không sửa page.", "Repository abstraction, REST contract, container stateless, PostgreSQL, LangGraph node hóa workflow."),
            ("Bảo trì", "Schema, dữ liệu và contract được kiểm soát phiên bản, kiểu dữ liệu chặt chẽ.", "TypeScript strict, Zod/Pydantic, SQLAlchemy models, Alembic migration, seed idempotent, lint/test/build."),
            ("AI an toàn", "Chỉ tư vấn bằng dữ liệu hệ thống; không tự tạo SKU, giá, chứng nhận hoặc claim môi trường.", "Structured retrieval + hybrid RAG + grounding; từ chối ngoài phạm vi; deterministic fallback."),
            ("Quan sát", "Có thể xác định mode retrieval và độ trễ mà không lưu câu hỏi rõ hoặc lộ secret.", "X-Request-ID, retrieval log dùng query hash, chunk/product IDs và endpoint health/rag."),
        ],
        widths_cm=[2.2, 6.2, 6.5],
        font_size=10,
    )

    writer.heading("3.3. Tác nhân và phân quyền")
    writer.heading("3.3.1. Danh sách tác nhân", level=3)
    writer.table(
        "Bảng 3.3. Tác nhân của hệ thống",
        ["Tác nhân", "Vai trò", "Phạm vi hiện tại"],
        [
            ("Guest", "Người truy cập chưa đăng nhập.", "Catalog, truy xuất, Advisor, chatbot, giỏ local và checkout COD demo."),
            ("Customer", "Người dùng có tài khoản và phiên hợp lệ.", "Toàn bộ chức năng Guest, duy trì phiên và đăng xuất; đơn demo có thể gắn email người nhận."),
            ("Admin", "Người vận hành catalog, đơn hàng và tri thức.", "Được phân tích trong kiến trúc mục tiêu; chưa mở UI/API quản trị ở bản demo."),
            ("AI Chatbot", "Tác nhân hệ thống tiếp nhận và trả lời hội thoại.", "LangGraph điều phối retrieval, grounding, Groq và fallback."),
            ("Groq API", "Dịch vụ LLM ngoài hệ thống.", "Chỉ nhận prompt và context đã giới hạn; không truy cập trực tiếp database."),
        ],
        widths_cm=[2.5, 5.0, 7.2],
    )
    writer.heading("3.3.2. Ma trận phân quyền", level=3)
    writer.table(
        "Bảng 3.4. Ma trận quyền theo tác nhân",
        ["Chức năng", "Guest", "Customer", "Admin"],
        [
            ("Xem/tìm/lọc sản phẩm", "Có", "Có", "Có"),
            ("Xem hồ sơ lô demo", "Có", "Có", "Có"),
            ("Dùng Advisor/chatbot", "Có", "Có", "Có"),
            ("Giỏ local và tạo đơn COD demo", "Có", "Có", "Có"),
            ("Quản lý phiên cá nhân", "Đăng ký/đăng nhập", "Xem phiên/đăng xuất", "Có"),
            ("Tạo/sửa/xóa sản phẩm", "Không", "Không", "Dự kiến"),
            ("Cập nhật trạng thái đơn", "Không", "Không", "Dự kiến"),
            ("Quản lý/re-index knowledge base", "Không", "Không", "Dự kiến"),
        ],
        widths_cm=[7.2, 2.5, 2.5, 2.5],
    )

    writer.page_break()
    writer.heading("3.4. Phân tích Use Case")
    writer.heading("3.4.1. Use Case tổng quát", level=3)
    writer.paragraph(
        "Biểu đồ tổng quát phân biệt rõ phần đã hiện thực và phần quản trị dự kiến. Customer "
        "kế thừa hành vi Guest; các use case Admin được thể hiện bằng nét đứt để giữ tính "
        "trung thực của báo cáo."
    )
    writer.figure(diagrams / "06-use-case-overview.png", "Hình 3.1. Biểu đồ Use Case tổng quát")

    writer.heading("3.4.2. Đặc tả các Use Case chính", level=3)
    use_cases = [
        (
            "Bảng 3.5. Đặc tả UC01 – Đăng nhập",
            [
                ("Actor", "Guest/Customer"),
                ("Tiền điều kiện", "Tài khoản tồn tại, chưa bị vô hiệu hóa; origin nằm trong allow-list."),
                ("Luồng chính", "1) Nhập email/mật khẩu. 2) Frontend validate. 3) Backend rate-limit và tìm email. 4) Argon2id xác minh mật khẩu. 5) Thu hồi/rotate phiên cũ. 6) Gửi cookie HttpOnly và thông tin user."),
                ("Luồng thay thế", "Email/mật khẩu sai trả 401; request vượt ngưỡng trả 429; origin sai trả 403; không tiết lộ email có tồn tại hay không."),
                ("Hậu điều kiện", "Phiên mới lưu bằng token hash trong PostgreSQL; trình duyệt không đọc được token."),
            ],
        ),
        (
            "Bảng 3.6. Đặc tả UC02 – Tìm kiếm và lọc sản phẩm",
            [
                ("Actor", "Guest/Customer"),
                ("Tiền điều kiện", "Catalog có sản phẩm published hoặc frontend đang dùng mock repository hợp lệ."),
                ("Luồng chính", "1) Nhập từ khóa/chọn filter. 2) Chuẩn hóa tiếng Việt. 3) Ghi search params lên URL. 4) Repository truy vấn nguồn dữ liệu. 5) Sắp xếp và render kết quả."),
                ("Luồng thay thế", "Không có kết quả: hiển thị empty state và nút xóa lọc; API lỗi: error boundary và thao tác thử lại; slug sai: not-found."),
                ("Hậu điều kiện", "URL có thể chia sẻ và tái tạo đúng trạng thái lọc; catalog gốc không bị thay đổi."),
            ],
        ),
        (
            "Bảng 3.7. Đặc tả UC03 – Tạo đơn hàng COD demo",
            [
                ("Actor", "Guest/Customer"),
                ("Tiền điều kiện", "Giỏ có ít nhất một variant còn hàng; người dùng chấp nhận điều kiện đơn demo."),
                ("Luồng chính", "1) Nhập người nhận. 2) React Hook Form/Zod validate. 3) Gửi item ID, quantity, grind và Idempotency-Key. 4) Backend đọc catalog, tính subtotal/phí. 5) Ghi order và item snapshots trong transaction. 6) Trả mã đơn."),
                ("Luồng thay thế", "Variant/kiểu xay sai hoặc hết hàng trả 422; key đã dùng trả lại đơn cũ; DB lỗi rollback toàn transaction."),
                ("Hậu điều kiện", "Đơn demo-confirmed được lưu; không có giao dịch thẻ hoặc gọi đơn vị vận chuyển."),
            ],
        ),
        (
            "Bảng 3.8. Đặc tả UC04 – Chatbot tư vấn",
            [
                ("Actor", "Guest/Customer, AI Chatbot"),
                ("Tiền điều kiện", "Knowledge base đã seed; khi vector bắt buộc, 21 chunks phải có embedding."),
                ("Luồng chính", "1) Gửi câu hỏi. 2) LangGraph phân loại intent. 3) Structured retrieval lấy ứng viên. 4) BM25 và pgvector tìm chunks. 5) RRF hợp nhất. 6) Grounding theo product published. 7) Groq sinh câu trả lời. 8) Trả message/actions."),
                ("Luồng thay thế", "Greeting hoặc ngoài phạm vi đi thẳng scope fallback; không có dữ liệu nói rõ giới hạn; Groq lỗi dùng deterministic fallback; vượt rate limit trả 429."),
                ("Hậu điều kiện", "Ghi retrieval log đã băm query; action chỉ trỏ tới route sản phẩm/truy xuất thật."),
            ],
        ),
        (
            "Bảng 3.9. Đặc tả UC05 – Admin quản lý sản phẩm (mở rộng)",
            [
                ("Actor", "Admin"),
                ("Tiền điều kiện", "Đã xác thực và có role ADMIN; chức năng thuộc phiên bản mở rộng."),
                ("Luồng chính", "1) Nhập/chỉnh catalog. 2) Backend kiểm tra role và schema. 3) Ghi transaction. 4) Tạo lại knowledge chunks/embedding liên quan. 5) Ghi audit trail."),
                ("Luồng thay thế", "Customer nhận 403; SKU/slug trùng nhận 409; dữ liệu sai nhận 422; re-index lỗi giữ phiên bản knowledge cũ."),
                ("Hậu điều kiện", "Catalog và chỉ mục tìm kiếm đồng bộ; thay đổi có người thực hiện và thời điểm."),
            ],
        ),
    ]
    for caption, rows in use_cases:
        writer.table(caption, ["Thuộc tính", "Đặc tả"], rows, widths_cm=[3.0, 11.8], font_size=10)

    writer.heading("3.5. Kiến trúc tổng thể hệ thống")
    writer.heading("3.5.1. Kiến trúc Frontend – Backend – Database", level=3)
    writer.paragraph(
        "Frontend Next.js App Router triển khai trên Vercel và giao tiếp qua REST. Rewrite "
        "/backend-api giữ request xác thực cùng origin. DuckDNS trỏ domain API đến VNPT "
        "Cloud; container minute_caddy kết thúc TLS và reverse proxy qua Docker network caddy "
        "đến dauvi-api:8000. Backend kết nối PostgreSQL qua mạng data nội bộ, không publish "
        "port database ra Internet."
    )
    writer.figure(diagrams / "07-system-architecture-rag.png", "Hình 3.2. Kiến trúc tổng thể và tích hợp AI")
    writer.heading("3.5.2. Kiến trúc tích hợp AI", level=3)
    writer.paragraph(
        "FastEmbed chạy trong backend để tạo embedding 384 chiều bằng mô hình multilingual "
        "MiniLM. Vector được lưu trực tiếp trong cột vector(384) của PostgreSQL và lập chỉ "
        "mục HNSW cosine. Groq là dịch vụ sinh ngôn ngữ, nhưng không truy cập database; nó "
        "chỉ nhận tối đa top-k chunks và tối đa ba sản phẩm đã qua grounding. Thiết kế này "
        "giữ retrieval và business rule ở phía hệ thống, thay vì giao toàn bộ quyết định cho LLM."
    )

    writer.heading("3.6. Thiết kế luồng xử lý")
    writer.heading("3.6.1. Đăng ký và đăng nhập", level=3)
    writer.paragraph(
        "Email được chuẩn hóa chữ thường và có unique constraint. Mật khẩu đạt chính sách độ "
        "dài mới được băm Argon2id. Sau đăng nhập, backend tạo token ngẫu nhiên, chỉ lưu SHA-256 "
        "hash trong user_sessions, còn giá trị token được gửi bằng cookie HttpOnly. Các thao tác "
        "thay đổi phiên kiểm tra Origin và giới hạn tần suất theo email/IP đã băm."
    )
    writer.heading("3.6.2. Tìm kiếm và lọc sản phẩm", level=3)
    writer.paragraph(
        "Bộ lọc species, region, process, roast, brew method, price và format được biểu diễn "
        "bằng search params. Hàm normalize loại dấu tiếng Việt và ký tự thừa trước khi so khớp. "
        "HTTP repository chuyển params thành query API; mock repository giữ cùng interface để "
        "test và phát triển không phụ thuộc backend."
    )
    writer.heading("3.6.3. Giỏ hàng và đặt hàng", level=3)
    writer.paragraph(
        "CartStore chỉ lưu productId, variantId, grind và quantity cần thiết để khôi phục giỏ. "
        "Khi checkout, backend không nhận unit_price từ client mà đọc giá hiện hành, kiểm tra "
        "quantity 1–99, format/grind và in_stock. order_items lưu snapshot tên, SKU, format, "
        "đơn giá và thành tiền để đơn cũ không bị thay đổi khi catalog cập nhật."
    )
    writer.heading("3.6.4. Quản trị sản phẩm", level=3)
    writer.paragraph(
        "Luồng quản trị được giữ ở mức thiết kế mở rộng: mọi endpoint /admin phải qua xác thực "
        "và role check ADMIN, cập nhật catalog trong transaction, sau đó phát sinh lại chunks "
        "và embedding liên quan. Không đưa API giả vào bản hiện tại chỉ để làm đầy phạm vi."
    )
    writer.heading("3.6.5. Chatbot tư vấn", level=3)
    writer.paragraph(
        "Luồng chatbot là graph hữu hạn có nhánh điều kiện. Intent greeting/out-of-scope không "
        "gọi retrieval/LLM không cần thiết. Các intent product, brew, traceability và commerce "
        "đi qua structured retrieval, hybrid retrieval, grounding, generate và audit. Mỗi node "
        "nhận/trả một phần AssistantGraphState, nhờ đó dễ kiểm thử và thay thế độc lập."
    )

    writer.page_break()
    writer.heading("3.7. Biểu đồ tuần tự")
    writer.heading("3.7.1. Sequence đăng nhập", level=3)
    writer.figure(diagrams / "08-login-sequence.png", "Hình 3.3. Sequence đăng ký và đăng nhập")
    writer.heading("3.7.2. Sequence đặt hàng", level=3)
    writer.figure(diagrams / "09-order-sequence.png", "Hình 3.4. Sequence tạo đơn hàng COD demo")
    writer.heading("3.7.3. Sequence chatbot", level=3)
    writer.figure(diagrams / "10-chatbot-sequence.png", "Hình 3.5. Sequence chatbot hybrid RAG")
    writer.heading("3.7.4. Sequence quản trị sản phẩm", level=3)
    writer.figure(diagrams / "11-admin-sequence-target.png", "Hình 3.6. Sequence quản trị sản phẩm – thiết kế mở rộng")

    writer.heading("3.8. Thiết kế cơ sở dữ liệu")
    writer.heading("3.8.1. ERD của schema hiện tại", level=3)
    writer.paragraph(
        "Schema được quản lý bằng Alembic. Migration 20260810_0004 bật extension vector, tạo "
        "knowledge_documents, knowledge_chunks, retrieval_logs và HNSW index. SQLite dùng JSON "
        "variant cho cột embedding trong unit test; production sử dụng kiểu vector(384)."
    )
    writer.figure(diagrams / "13-current-erd-pgvector.png", "Hình 3.7. ERD rút gọn của PostgreSQL và pgvector")
    writer.heading("3.8.2. Các bảng dữ liệu", level=3)
    writer.table(
        "Bảng 3.10. Nhóm bảng đã triển khai",
        ["Nhóm", "Bảng", "Mục đích"],
        [
            ("Catalog", "products, product_variants", "Sản phẩm, taxonomy, flavor profile, format, grind, giá và trạng thái còn hàng."),
            ("Truy xuất", "coffee_lots, evidence_items, lot_timeline_events", "Hồ sơ lô, evidence level, nguồn và timeline sáu bước."),
            ("Đơn hàng", "orders, order_items", "Người nhận, tổng tiền server-side, idempotency và snapshot dòng hàng."),
            ("Xác thực", "users, user_sessions, auth_attempts", "Tài khoản, mật khẩu băm, phiên băm và rate-limit audit."),
            ("Chatbot", "assistant_requests", "Giới hạn tần suất theo client hash."),
            ("RAG", "knowledge_documents, knowledge_chunks", "Nguồn tri thức, chunks, metadata và embedding vector(384)."),
            ("Quan sát RAG", "retrieval_logs", "Query hash, intent, IDs kết quả, latency và cờ vector/LLM."),
        ],
        widths_cm=[2.2, 5.2, 7.4],
    )
    writer.heading("3.8.3. Quan hệ và ràng buộc dữ liệu", level=3)
    writer.bullet("Catalog: ", "một product có nhiều variants, lots và knowledge chunks; slug/SKU là duy nhất.")
    writer.bullet("Traceability: ", "một lot có nhiều evidence items và đúng sáu timeline events theo stage duy nhất.")
    writer.bullet("Orders: ", "một order có nhiều items; order_code và idempotency_key là duy nhất; total bằng subtotal + shipping_fee.")
    writer.bullet("Authentication: ", "email duy nhất; session chỉ liên kết user hợp lệ và có expires_at/revoked_at.")
    writer.bullet("RAG: ", "một document có nhiều chunks; (document_id, chunk_index) duy nhất; chunk có thể liên kết product hoặc lot thật.")
    writer.heading("3.8.4. Seed data", level=3)
    writer.table(
        "Bảng 3.11. Sáu sản phẩm và mã lô seed",
        ["Sản phẩm", "Vùng", "Sơ chế", "Giá 250 g", "Mã lô"],
        [
            ("TRS1 Tây Nguyên Daily Phin", "Gia Lai", "Natural", "99.000 ₫", "TRS1-GL-26-N01"),
            ("TR4 Đắk Lắk Traceable Robusta", "Đắk Lắk", "Natural", "119.000 ₫", "TR4-DLK-26-N02"),
            ("TR9 Large Bean Fine Robusta", "Đắk Lắk", "Honey", "139.000 ₫", "TR9-DLK-26-H01"),
            ("Xanh Lùn TS5 Bảo Lâm Honey", "Bảo Lâm", "Honey", "159.000 ₫", "XLTS5-BL-26-H01"),
            ("Catimor Đà Lạt Washed", "Đà Lạt", "Washed", "139.000 ₫", "CAT-DL-26-W01"),
            ("Bourbon Langbiang Honey", "Langbiang", "Honey", "199.000 ₫", "BBN-LB-26-H01"),
        ],
        widths_cm=[5.2, 2.3, 2.0, 2.0, 3.3],
        font_size=9.8,
    )
    writer.paragraph(
        "Seed RAG gồm 7 knowledge documents và 21 chunks: mỗi sản phẩm có ba chunks về hồ "
        "sơ hương vị, cách pha/quy cách và truy xuất; tài liệu chính sách chung có ba chunks "
        "về minh bạch, checkout demo và phạm vi chatbot. Seed idempotent cập nhật content_hash "
        "và chỉ tạo lại embedding khi nội dung hoặc embedding model thay đổi."
    )

    writer.page_break()
    writer.heading("3.9. Thiết kế REST API")
    writer.paragraph(
        "API sử dụng JSON, số tiền nguyên VND, ngày ISO và lỗi chuẩn hóa dạng message/code. "
        "Các endpoint nghiệp vụ có prefix /api/v1; healthcheck nằm ở root để Caddy và Docker "
        "kiểm tra không phụ thuộc prefix."
    )
    writer.table(
        "Bảng 3.12. Các endpoint REST đã triển khai",
        ["Method", "Endpoint", "Chức năng", "Quyền"],
        [
            ("GET", "/api/v1/products", "Danh sách, tìm kiếm, lọc và sort", "Public"),
            ("GET", "/api/v1/products/featured", "Sản phẩm nổi bật", "Public"),
            ("GET", "/api/v1/products/{slug}", "Chi tiết theo slug", "Public"),
            ("GET", "/api/v1/lots/featured", "Các lô nổi bật", "Public"),
            ("GET", "/api/v1/lots/{lotCode}", "Passport truy xuất", "Public"),
            ("POST", "/api/v1/advisor/recommendations", "Top 3 theo bộ trọng số", "Public"),
            ("POST", "/api/v1/assistant/messages", "Chatbot LangGraph hybrid RAG", "Public/rate-limit"),
            ("POST", "/api/v1/auth/register", "Đăng ký và tạo phiên", "Public"),
            ("POST", "/api/v1/auth/login", "Đăng nhập và rotate phiên", "Public"),
            ("GET", "/api/v1/auth/session", "Đọc user của phiên", "Session"),
            ("POST", "/api/v1/auth/logout", "Thu hồi phiên", "Session"),
            ("POST", "/api/v1/orders", "Tạo đơn COD demo", "Public + idempotency"),
            ("GET", "/health/live", "Tiến trình còn sống", "Infrastructure"),
            ("GET", "/health/ready", "DB/vector bắt buộc sẵn sàng", "Infrastructure"),
            ("GET", "/health/rag", "Workflow, retrieval mode và số vectors", "Infrastructure"),
        ],
        widths_cm=[1.5, 5.3, 5.6, 2.4],
        font_size=9.5,
    )

    writer.heading("3.10. Thiết kế Chatbot AI")
    writer.heading("3.10.1. Workflow LangGraph", level=3)
    writer.figure(diagrams / "12-langgraph-workflow.png", "Hình 3.8. Workflow LangGraph và cơ chế grounding")
    writer.paragraph(
        "StateGraph được biên dịch với các node understand, structured_retrieval, "
        "hybrid_retrieval, grounding, generate, scope_fallback và audit. Edge điều kiện sau "
        "understand tách greeting/out-of-scope khỏi luồng RAG. Graph trả đúng contract "
        "AssistantResponse gồm message và actions, không trả raw prompt hoặc vector."
    )
    writer.heading("3.10.2. Phân tích nhu cầu người dùng", level=3)
    writer.paragraph(
        "Node understand chuẩn hóa Unicode, bỏ dấu để so khớp ổn định và phân loại bảy intent: "
        "greeting, product-advice, product-fact, traceability, brew-guide, commerce và "
        "out-of-scope. Structured retrieval đồng thời nhận diện ngân sách, từ khóa giống cà "
        "phê, cách pha và mã sản phẩm để tạo tập ứng viên ban đầu."
    )
    writer.heading("3.10.3. Structured Product Retrieval", level=3)
    writer.paragraph(
        "Truy vấn có cấu trúc sử dụng product/variant/lots đã published. Với câu có ngân sách, "
        "hard filter loại variant vượt mức trước khi đưa product ID vào hybrid retrieval. Cách "
        "này bảo đảm yêu cầu “dưới 120.000 ₫” không bị LLM phá vỡ và action luôn có route thật."
    )
    writer.heading("3.10.4. BM25 và Vector Retrieval", level=3)
    writer.paragraph(
        "BM25 được cài đặt trực tiếp trên title + content đã tokenize; k1 = 1,5 và b = 0,75. "
        "Nhánh vector dùng FastEmbed tạo embedding query 384 chiều, sau đó SQLAlchemy/pgvector "
        "xếp hạng theo cosine distance. HNSW giảm chi phí tìm kiếm khi knowledge base mở rộng. "
        "Nếu môi trường phát triển dùng SQLite hoặc embedding lỗi và vector không bắt buộc, "
        "BM25 vẫn cung cấp retrieval có thể kiểm thử."
    )
    writer.heading("3.10.5. Ranking và Grounding", level=3)
    writer.paragraph(
        "Hai danh sách được hợp nhất bằng Reciprocal Rank Fusion: RRF(d) = Σ 1/(k + rankᵢ(d)), "
        "với k = 60. Top 6 chunks được tải làm knowledge context. Node grounding chỉ chấp nhận "
        "product ID tồn tại và published, giữ tối đa ba sản phẩm. Prompt Groq bao gồm catalog "
        "bounded và retrieved knowledge; actions được backend dựng riêng thay vì tin output LLM."
    )
    writer.heading("3.10.6. Fallback và chống hallucination", level=3)
    writer.table(
        "Bảng 3.13. Các tình huống phản hồi chatbot",
        ["Tình huống", "Xử lý", "Kết quả"],
        [
            ("Có sản phẩm phù hợp", "Hybrid retrieval + grounding + Groq.", "Gợi ý tối đa 3 sản phẩm, lý do và action thật."),
            ("Không khớp hoàn toàn", "Giữ hard constraint; dùng ứng viên gần nhất chỉ khi không vi phạm ngân sách/format.", "Nói rõ mức độ phù hợp, không tuyên bố khớp tuyệt đối."),
            ("Không có dữ liệu", "Scope fallback hoặc empty retrieval.", "Nói chưa có thông tin và giới hạn phạm vi."),
            ("Groq lỗi/timeout", "Deterministic fallback từ grounded products.", "Endpoint vẫn trả câu trả lời catalog, không bịa."),
            ("Ngoài phạm vi", "Không gọi LLM; trả lời từ chối có hướng dẫn.", "Không trả kiến thức tự do ngoài hệ thống."),
        ],
        widths_cm=[3.0, 6.2, 5.6],
    )

    writer.heading("3.11. Thiết kế xác thực, phân quyền và bảo mật")
    writer.paragraph(
        "Phiên bản hiện tại có hai mức truy cập thực thi: public và authenticated session. "
        "Role ADMIN là thiết kế mở rộng; khi triển khai phải kiểm tra role ở backend, không chỉ "
        "ẩn nút trên frontend. Hệ thống không dùng JWT/localStorage cho đăng nhập mà dùng session "
        "opaque, phù hợp với rewrite same-origin của Vercel."
    )
    writer.table(
        "Bảng 3.14. Business rules và biện pháp bảo vệ",
        ["Rủi ro/Rule", "Biện pháp thiết kế"],
        [
            ("Email trùng", "Unique constraint và trả lỗi xung đột không làm lộ dữ liệu nhạy cảm."),
            ("Lộ mật khẩu", "Argon2id; không log password; secret production tối thiểu 32 ký tự."),
            ("Đánh cắp session", "Token ngẫu nhiên chỉ gửi HttpOnly/Secure/SameSite; DB lưu token_hash; hỗ trợ revoke/expiry."),
            ("CSRF/CORS", "Same-origin rewrite, Origin validation cho auth, CORS origin allow-list và credentials có kiểm soát."),
            ("Brute force/chat abuse", "Rate limit theo cửa sổ thời gian với client/email/IP đã băm."),
            ("Giả giá/quantity", "Server đọc unit price; quantity 1–99; validate grind/in_stock; transaction và idempotency."),
            ("AI hallucination", "Chỉ product/lot ID thật; bounded context; fallback; không tạo claim/chứng nhận/review."),
            ("Lộ hạ tầng", "DB ở Docker network internal; Caddy là cổng TLS; security headers và TrustedHost."),
            ("Secret bị public", "POSTGRES_PASSWORD, SESSION_SECRET và GROQ_API_KEY chỉ nằm trong docker/.env trên VM."),
        ],
        widths_cm=[4.2, 10.6],
        font_size=10,
    )


def write_chapter_four(writer: Writer) -> None:
    writer.paragraph(
        "Chương này trình bày kết quả cài đặt, cách chạy, kiểm thử và triển khai hệ thống từ "
        "repository thực tế. Các lệnh sử dụng đúng cấu trúc frontend/, backend/ và docker/; "
        "không yêu cầu API giả trong Next.js."
    )
    writer.heading("4.1. Môi trường và cấu trúc cài đặt")
    writer.heading("4.1.1. Công nghệ cốt lõi", level=3)
    writer.table(
        "Bảng 4.1. Công nghệ cốt lõi đã sử dụng",
        ["Lớp", "Công nghệ", "Vai trò"],
        [
            ("Frontend", "Next.js 16.2.10, React 19.2.7, TypeScript 5.9", "App Router, RSC mặc định, metadata/JSON-LD, UI tương tác có type."),
            ("UI", "Tailwind CSS 4, Radix/shadcn pattern, Motion, Lucide", "Design system editorial, responsive, dialog, animation và icon."),
            ("State/Form", "Zustand 5, React Hook Form, Zod 4", "Cart persistence, form state và validation."),
            ("Backend", "Python 3.12, FastAPI 0.139.2, Pydantic Settings", "REST API, validation, middleware và cấu hình."),
            ("Data", "SQLAlchemy 2.0.51, Alembic 1.18.5, PostgreSQL 17", "ORM, migration, transaction và lưu trữ bền vững."),
            ("AI/RAG", "LangGraph 1.2.10, FastEmbed 0.8.0, pgvector 0.5.0, Groq", "Graph orchestration, embedding, vector search và grounded generation."),
            ("Kiểm thử", "Vitest, Testing Library, Playwright, Pytest, Ruff", "Unit/UI/E2E và kiểm tra chất lượng Python."),
            ("Triển khai", "Vercel, Docker Compose, VNPT Cloud, Caddy, DuckDNS", "Frontend edge, container backend/database, TLS và reverse proxy."),
        ],
        widths_cm=[2.0, 5.8, 7.0],
        font_size=9.8,
    )
    writer.heading("4.1.2. Cấu trúc repository", level=3)
    writer.code(
        "dauvi.coffee/\n"
        "├── frontend/          # Next.js App Router, UI, repositories, tests\n"
        "├── backend/           # FastAPI, models, migrations, seed, LangGraph RAG\n"
        "├── docker/            # Dockerfile, Compose, Caddy network override\n"
        "└── docs/              # Đặc tả, API contract, deployment và báo cáo"
    )
    writer.paragraph(
        "Cách đóng gói này cho phép Vercel chọn frontend làm Root Directory, trong khi VNPT "
        "Cloud build backend từ context root để Dockerfile truy cập backend/ và docker/."
    )

    writer.heading("4.2. Cài đặt frontend")
    writer.heading("4.2.1. Kiến trúc và data access", level=3)
    writer.paragraph(
        "Next.js dùng Server Components mặc định cho page đọc dữ liệu; Client Components chỉ "
        "xuất hiện ở bộ lọc, selector, cart, form, map tương tác, carousel và chatbot. Factory "
        "createRepositories chọn mock hoặc HTTP theo NEXT_PUBLIC_DATA_SOURCE. API client hỗ trợ "
        "base URL public và rewrite /backend-api, nên page không phụ thuộc trực tiếp mock JSON."
    )
    writer.heading("4.2.2. Các route và giao diện đã hoàn thành", level=3)
    writer.table(
        "Bảng 4.2. Route frontend chính",
        ["Route", "Nội dung"],
        [
            ("/", "Homepage editorial, 4 discovery cards, flavor map, featured, traceability carousel, Advisor CTA."),
            ("/shop", "Search/filter/sort bằng URL params, active chips và empty state."),
            ("/shop/[slug]", "6 trang sản phẩm, variant/grind/quantity, passport, related và sticky CTA mobile."),
            ("/traceability", "Lookup uppercase, demo chips, evidence levels và featured lots."),
            ("/traceability/[lotCode]", "Passport, timeline 6 bước, evidence, disclosure và liên kết sản phẩm."),
            ("/advisor", "Quiz 6 bước và top 3 recommendation có score/reasons."),
            ("/cart, /checkout", "Cart persist, free-shipping progress và form COD demo."),
            ("/login, /register", "Form tài khoản thật nối backend session."),
            ("/story, /brew-guide", "Nội dung thương hiệu và hướng dẫn pha tại nhà."),
            ("System routes", "loading, error, global-error, not-found, robots.txt và sitemap.xml."),
        ],
        widths_cm=[4.0, 10.8],
        font_size=10,
    )
    writer.heading("4.2.3. Responsive, accessibility và SEO", level=3)
    writer.paragraph(
        "Layout đã được kiểm tra ở desktop và các breakpoint quy định. Header desktop/mobile, "
        "bottom navigation và sticky add-to-cart có khoảng đệm tránh che nội dung. Dialog dùng "
        "Radix focus trap; focus ring, reduced-motion, alt text bản đồ/ảnh và flavor label được "
        "đưa vào component. Metadata App Router đi kèm Organization, Product và Breadcrumb JSON-LD."
    )

    writer.heading("4.3. Cài đặt backend và cơ sở dữ liệu")
    writer.heading("4.3.1. FastAPI và middleware", level=3)
    writer.paragraph(
        "FastAPI đăng ký router products, lots, advisor, orders, auth và assistant dưới /api/v1. "
        "Middleware gồm GZip, TrustedHost, CORS allow-list, request ID và security headers. "
        "Validation error và HTTP error được chuyển về body message/code thống nhất để frontend "
        "hiển thị trạng thái lỗi mà không lộ stack trace production."
    )
    writer.heading("4.3.2. Xác thực và đơn hàng", level=3)
    writer.paragraph(
        "Auth service dùng Argon2id và session opaque. Endpoint register/login tạo cookie; "
        "session/logout đọc và thu hồi phiên ở PostgreSQL. Order service nhận Idempotency-Key, "
        "đọc product/variant hiện hành, kiểm tra grind/in_stock/quantity, tính miễn phí giao hàng "
        "từ 499.000 ₫ và ghi order + item snapshots trong transaction."
    )
    writer.heading("4.3.3. LangGraph hybrid RAG", level=3)
    writer.paragraph(
        "Backend đã cài LangGraph thật, không chỉ mô tả trong tài liệu. Workflow có bảy node và "
        "nhánh điều kiện; graph được compile một lần và nhận session/settings qua runtime context "
        "riêng cho từng request. Retrieval kết hợp BM25 với pgvector cosine, sau đó RRF. FastEmbed model "
        "được tải sẵn trong Docker image để runtime không phụ thuộc tải model. Seed tạo 21 vectors "
        "khi VECTOR_SEARCH_ENABLED=true. health/rag xác nhận graph, retrieval mode, số chunks/vector "
        "và provider LLM."
    )
    writer.heading("4.3.4. Migration và seed", level=3)
    writer.paragraph(
        "Entrypoint container chạy alembic upgrade head rồi python -m app.seed trước Uvicorn. "
        "Migration mới dùng CREATE EXTENSION IF NOT EXISTS vector và tạo HNSW index cosine. "
        "Seed idempotent cập nhật sáu sản phẩm, sáu lô, evidence/timeline và knowledge base; do "
        "đó restart container không tạo bản ghi catalog trùng."
    )

    writer.page_break()
    writer.heading("4.4. Chạy toàn bộ repository trên máy phát triển")
    writer.heading("4.4.1. Backend", level=3)
    writer.paragraph(
        "Cách nhanh dùng SQLite và BM25; đặt VECTOR_SEARCH_ENABLED=false trong backend/.env."
    )
    writer.code(
        "cd backend\n"
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "pip install -r requirements-dev.txt\n"
        "Copy-Item .env.example .env\n"
        "alembic upgrade head\n"
        "python -m app.seed\n"
        "uvicorn app.main:app --reload"
    )
    writer.paragraph(
        "API chạy tại http://localhost:8000; OpenAPI ở /docs; kiểm tra /health/ready và "
        "/health/rag. Muốn kiểm thử vector giống production, dùng Docker Compose với image "
        "pgvector/pgvector:0.8.6-pg17-bookworm."
    )
    writer.heading("4.4.2. Frontend", level=3)
    writer.code(
        "cd frontend\n"
        "Copy-Item .env.example .env.local\n"
        "pnpm install --frozen-lockfile\n"
        "pnpm dev"
    )
    writer.paragraph(
        "Đặt NEXT_PUBLIC_DATA_SOURCE=http và NEXT_PUBLIC_API_BASE_URL=http://localhost:8000/api/v1 "
        "để dùng backend local. Mở http://localhost:3000. Nếu chỉ kiểm tra UI độc lập, giữ "
        "NEXT_PUBLIC_DATA_SOURCE=mock."
    )
    writer.heading("4.4.3. Các lệnh kiểm tra chất lượng", level=3)
    writer.code(
        "# Frontend\n"
        "cd frontend\n"
        "pnpm lint\n"
        "pnpm typecheck\n"
        "pnpm test\n"
        "pnpm build\n"
        "pnpm test:e2e\n\n"
        "# Backend\n"
        "cd ../backend\n"
        "python -m ruff check .\n"
        "python -m pytest tests -q"
    )

    writer.heading("4.5. Đóng gói Docker")
    writer.paragraph(
        "Dockerfile dùng Python 3.12 slim, user không đặc quyền UID 10001, root filesystem "
        "read-only, tmpfs /tmp, healthcheck và không ghi pip cache. Embedding model được preload "
        "ở build stage. Compose tách app/data network; database chỉ ở network internal, còn "
        "backend được nối thêm network caddy bằng compose.caddy.yml."
    )
    writer.code(
        "cd /opt/dauvi.coffee\n"
        "docker pull pgvector/pgvector:0.8.6-pg17-bookworm\n"
        "docker build --network=host --pull \\\n+  --build-arg EMBEDDING_MODEL=sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2 \\\n+  -f docker/backend.Dockerfile -t dau-vi-backend:latest .\n"
        "docker compose --env-file docker/.env \\\n+  -f docker/compose.yml -f docker/compose.caddy.yml \\\n+  up -d --no-build database backend"
    )
    writer.paragraph(
        "Tùy chọn --network=host được dùng khi Docker daemon trên cloud gặp lỗi phân giải DNS "
        "trong build, dù host vẫn truy cập PyPI/Hugging Face bình thường."
    )

    writer.heading("4.6. Triển khai production")
    writer.heading("4.6.1. Frontend trên Vercel", level=3)
    writer.paragraph(
        "Kết nối repository GitHub và đặt Root Directory là frontend. Các biến Production cần "
        "thiết lập như sau; Groq key và database secret không được đưa lên Vercel."
    )
    writer.code(
        "NEXT_PUBLIC_SITE_URL=https://dauvi-coffee.vercel.app\n"
        "NEXT_PUBLIC_API_BASE_URL=https://dauvi-coffee.vercel.app/backend-api\n"
        "API_BASE_URL=https://dauvi-api.duckdns.org/api/v1\n"
        "BACKEND_PROXY_ORIGIN=https://dauvi-api.duckdns.org\n"
        "NEXT_PUBLIC_DATA_SOURCE=http\n"
        "NEXT_PUBLIC_ENABLE_CHECKOUT=true\n"
        "NEXT_PUBLIC_ENABLE_AUTH=true\n"
        "NEXT_PUBLIC_ENABLE_CHATBOT_API=true"
    )
    writer.heading("4.6.2. Backend trên VNPT Cloud", level=3)
    writer.paragraph(
        "Repository đặt tại /opt/dauvi.coffee. Trước khi đổi image PostgreSQL chuẩn sang image "
        "cùng major có pgvector, cần sao lưu volume bằng pg_dump."
    )
    writer.code(
        "cd /opt/dauvi.coffee\n"
        "git pull --ff-only\n"
        "mkdir -p /opt/dauvi-backups\n"
        "docker compose --env-file docker/.env -f docker/compose.yml exec -T database \\\n+  pg_dump -U dauvi -d dauvi -Fc \\\n+  > \"/opt/dauvi-backups/dauvi-pre-rag-$(date +%F-%H%M).dump\""
    )
    writer.paragraph(
        "Sau đó thực hiện lệnh build ở Mục 4.5. Entrypoint tự migration/seed. Các biến bắt buộc "
        "trong docker/.env gồm POSTGRES_PASSWORD, SESSION_SECRET, CORS_ORIGINS, RAG/VECTOR flags, "
        "GROQ_API_KEY, GROQ_MODEL, BACKEND_PORT=18081 và CADDY_NETWORK=caddy."
    )
    writer.heading("4.6.3. DuckDNS, Caddy và HTTPS", level=3)
    writer.code(
        "# Caddyfile\n"
        "dauvi-api.duckdns.org {\n"
        "  encode zstd gzip\n"
        "  reverse_proxy dauvi-api:8000\n"
        "}\n\n"
        "docker network inspect caddy >/dev/null 2>&1 || docker network create caddy\n"
        "docker network inspect caddy --format '{{json .Containers}}' | grep -q 'minute_caddy' \\\n+  || docker network connect caddy minute_caddy\n"
        "docker exec minute_caddy caddy validate --config /etc/caddy/Caddyfile\n"
        "docker exec minute_caddy caddy reload --config /etc/caddy/Caddyfile"
    )
    writer.paragraph(
        "Caddy gọi tên service qua network caddy, không gọi 127.0.0.1 trong container. Port "
        "18081 chỉ phục vụ kiểm tra local trên VM vì host port 8000 đã thuộc service khác."
    )

    writer.page_break()
    writer.heading("4.7. Kiểm tra sau triển khai")
    writer.heading("4.7.1. Healthcheck và RAG", level=3)
    writer.code(
        "curl --max-time 15 -fsS http://127.0.0.1:18081/health/ready; echo\n"
        "curl --max-time 15 -fsS https://dauvi-api.duckdns.org/health/rag; echo\n"
        "docker compose --env-file docker/.env -f docker/compose.yml exec -T database \\\n+  psql -U dauvi -d dauvi -c \\\n+  'SELECT count(*) AS chunks, count(embedding) AS vectors FROM knowledge_chunks;'"
    )
    writer.paragraph(
        "Kết quả production mong đợi là status=ready, workflow=langgraph, "
        "retrieval=bm25+pgvector và 21/21 chunks có embedding."
    )
    writer.heading("4.7.2. Kiểm thử chatbot", level=3)
    writer.code(
        "curl --max-time 30 -fsS -X POST \\\n+  https://dauvi-api.duckdns.org/api/v1/assistant/messages \\\n+  -H 'Content-Type: application/json' \\\n+  -d '{\"message\":\"Tư vấn cà phê pha phin đậm dưới 120.000 đồng\"}'"
    )
    writer.paragraph(
        "Ca kiểm tra phải chỉ gợi ý TRS1 hoặc TR4 và không đưa sản phẩm vượt ngân sách. Các ca "
        "bổ sung gồm pour-over ít đắng (Catimor/Bourbon), drip bag (Catimor), mã lô "
        "TR4-DLK-26-N02 và câu hỏi ngoài phạm vi để xác nhận fallback."
    )
    writer.heading("4.7.3. Kết quả kiểm tra mã nguồn", level=3)
    writer.table(
        "Bảng 4.3. Kết quả kiểm tra tại thời điểm cập nhật báo cáo",
        ["Hạng mục", "Kết quả"],
        [
            ("Backend Ruff", "Đạt – không còn lỗi lint."),
            ("Backend Pytest", "Đạt – 11 test, gồm health/rag, LangGraph, BM25/RRF, retrieval log và out-of-scope fallback."),
            ("Docker Compose config", "Đạt – compose.yml + compose.caddy.yml hợp lệ với env mẫu."),
            ("Frontend ESLint/TypeScript", "Đạt – lint và typecheck không có lỗi."),
            ("Frontend Vitest", "Đạt – 24/24 unit test."),
            ("Frontend Playwright", "Đạt – 13/13 E2E, gồm cart persist, checkout, chatbot, traceability và overflow ở 4 viewport."),
            ("Frontend route audit production", "17 route nghiệp vụ trả HTTP 200; không có console error, failed request hoặc overflow ngang ở viewport 1440 px."),
            ("Frontend build", "Đạt – Next.js production build sinh 15 route pattern App Router; các slug động dùng chung route [slug]/[lotCode]."),
        ],
        widths_cm=[5.0, 9.8],
    )

    writer.heading("4.8. Giới hạn hiện tại và hướng hoàn thiện")
    writer.bullet("Quản trị: ", "chưa có Admin UI/API/RBAC; hiện mới có thiết kế mở rộng và phải triển khai trước khi vận hành thương mại.")
    writer.bullet("Đơn hàng: ", "là luồng COD demo, chưa tích hợp thanh toán, vận chuyển hoặc quản lý tồn kho theo số lượng.")
    writer.bullet("Dữ liệu truy xuất: ", "farm/cooperative/lot là Demo Data và luôn có disclosure; chưa kết nối nhà cung cấp xác minh thật.")
    writer.bullet("RAG: ", "knowledge base hiện giới hạn 6 sản phẩm/21 chunks; chưa có CMS, lịch re-index nền hoặc đánh giá retrieval quy mô lớn.")
    writer.bullet("Chatbot: ", "chưa streaming token và chưa lưu lịch sử hội thoại dài hạn; ưu tiên hiện tại là grounding và bảo vệ phạm vi.")
    writer.bullet("Vận hành: ", "cần bổ sung monitoring/alert, backup định kỳ, restore drill và secret manager khi chuyển khỏi quy mô đồ án.")

    writer.heading("4.9. Quy trình cập nhật và vận hành lại")
    writer.paragraph(
        "Mỗi lần cập nhật backend cần pull fast-forward, build image mới, chạy Compose và kiểm "
        "tra health trước khi xóa image cũ. Migration luôn đi kèm bản sao lưu PostgreSQL."
    )
    writer.code(
        "cd /opt/dauvi.coffee\n"
        "git pull --ff-only\n"
        "docker build --network=host --pull \\\n+  --build-arg EMBEDDING_MODEL=sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2 \\\n+  -f docker/backend.Dockerfile -t dau-vi-backend:latest .\n"
        "docker compose --env-file docker/.env \\\n+  -f docker/compose.yml -f docker/compose.caddy.yml \\\n+  up -d --no-build backend\n"
        "curl --max-time 15 -fsS https://dauvi-api.duckdns.org/health/rag; echo\n"
        "docker image prune -f"
    )
    writer.paragraph(
        "Tài liệu lệnh đầy đủ được lưu tại docs/deployment.md. Frontend được redeploy từ Vercel "
        "sau khi push Git hoặc sau khi thay biến môi trường. Trình tự khôi phục gồm dừng backend, "
        "khôi phục bản pg_dump tương ứng, chạy lại migration của phiên bản code được chọn và kiểm "
        "tra ba health endpoint trước khi mở traffic."
    )


def enable_field_updates(document: DocumentObject) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")
    for instruction in document.element.body.xpath(".//w:instrText"):
        value = instruction.text or ""
        if "TOC" not in value:
            continue
        if '\\o "1-2"' in value:
            instruction.text = ' TOC \\o "1-2" \\h \\z \\t "CAP3,3" '
        elif '"hình,1"' in value:
            instruction.text = ' TOC \\h \\z \\t "VN_Caption,1" '
        elif '"bảng,1"' in value:
            instruction.text = ' TOC \\h \\z \\t "VN_TableCaption,1" '


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("document", type=Path)
    parser.add_argument("--diagrams", type=Path, default=Path("docs/report-assets/diagrams"))
    args = parser.parse_args()

    target = args.document.resolve()
    diagrams = args.diagrams.resolve()
    backup_dir = target.parent / "backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    backup = backup_dir / f"{target.stem}.before-ch3-ch4.docx"
    if not backup.exists():
        shutil.copy2(target, backup)

    document = Document(target)
    chapter_three = find_paragraph(document, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
    chapter_four = find_paragraph(document, "CHƯƠNG 4. CÀI ĐẶT CHƯƠNG TRÌNH VÀ TRIỂN KHAI")
    conclusion = find_paragraph(document, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

    clear_between(chapter_three, chapter_four)
    clear_between(chapter_four, conclusion)
    write_chapter_three(Writer(document, chapter_four), diagrams)
    write_chapter_four(Writer(document, conclusion))
    enable_field_updates(document)
    document.core_properties.modified = datetime.now()

    temporary = target.with_suffix(".tmp.docx")
    document.save(temporary)
    os.replace(temporary, target)
    print(target)


if __name__ == "__main__":
    main()
