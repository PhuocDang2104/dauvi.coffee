from __future__ import annotations

import argparse
import collections
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _length_cm(value: object) -> float | None:
    if value is None:
        return None
    return round(value.cm, 2)


def _length_pt(value: object) -> float | None:
    if value is None:
        return None
    return round(value.pt, 2)


def _iter_blocks(document: DocumentObject):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def inspect_report(path: Path) -> dict[str, object]:
    document = Document(path)
    style_counts = collections.Counter(
        paragraph.style.name for paragraph in document.paragraphs if paragraph.text.strip()
    )
    styles = {}
    for style_name in style_counts:
        style = document.styles[style_name]
        paragraph_format = style.paragraph_format
        styles[style_name] = {
            "base_style": style.base_style.name if style.base_style else None,
            "font": style.font.name,
            "font_size_pt": _length_pt(style.font.size),
            "bold": style.font.bold,
            "italic": style.font.italic,
            "alignment": (
                style.paragraph_format.alignment.name
                if style.paragraph_format.alignment is not None
                else None
            ),
            "line_spacing": str(paragraph_format.line_spacing),
            "space_before_pt": _length_pt(paragraph_format.space_before),
            "space_after_pt": _length_pt(paragraph_format.space_after),
            "keep_with_next": paragraph_format.keep_with_next,
            "page_break_before": paragraph_format.page_break_before,
        }
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        paragraph_properties = paragraph._p.pPr
        outline_level = None
        if paragraph_properties is not None:
            outline = paragraph_properties.find(qn("w:outlineLvl"))
            if outline is not None:
                outline_level = outline.get(qn("w:val"))
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name,
                "outline_level": outline_level,
                "text": text,
            }
        )

    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([" ".join(cell.text.split()) for cell in row.cells])
        tables.append(
            {
                "index": table_index,
                "style": table.style.name if table.style else None,
                "rows": rows,
            }
        )

    paragraph_indexes = {paragraph._p: index for index, paragraph in enumerate(document.paragraphs)}
    table_indexes = {table._tbl: index for index, table in enumerate(document.tables)}
    ordered_blocks = []
    for block in _iter_blocks(document):
        if isinstance(block, Paragraph):
            image_targets = []
            for blip in block._p.xpath(".//a:blip"):
                relationship_id = blip.get(qn("r:embed"))
                if relationship_id and relationship_id in document.part.rels:
                    image_targets.append(document.part.rels[relationship_id].target_ref)
            text = " ".join(block.text.split())
            if text or image_targets:
                ordered_blocks.append(
                    {
                        "type": "paragraph",
                        "index": paragraph_indexes.get(block._p),
                        "style": block.style.name,
                        "text": text,
                        "images": image_targets,
                    }
                )
        else:
            first_row = [" ".join(cell.text.split()) for cell in block.rows[0].cells]
            ordered_blocks.append(
                {
                    "type": "table",
                    "index": table_indexes.get(block._tbl),
                    "style": block.style.name if block.style else None,
                    "rows": len(block.rows),
                    "columns": len(block.columns),
                    "first_row": first_row,
                }
            )

    sections = []
    for index, section in enumerate(document.sections):
        sections.append(
            {
                "index": index,
                "page_width_cm": _length_cm(section.page_width),
                "page_height_cm": _length_cm(section.page_height),
                "top_margin_cm": _length_cm(section.top_margin),
                "bottom_margin_cm": _length_cm(section.bottom_margin),
                "left_margin_cm": _length_cm(section.left_margin),
                "right_margin_cm": _length_cm(section.right_margin),
                "header_distance_cm": _length_cm(section.header_distance),
                "footer_distance_cm": _length_cm(section.footer_distance),
                "different_first_page": section.different_first_page_header_footer,
                "header": [p.text for p in section.header.paragraphs if p.text.strip()],
                "first_page_header": [
                    p.text for p in section.first_page_header.paragraphs if p.text.strip()
                ],
                "footer": [p.text for p in section.footer.paragraphs if p.text.strip()],
                "first_page_footer": [
                    p.text for p in section.first_page_footer.paragraphs if p.text.strip()
                ],
            }
        )

    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        media = []
        for name in names:
            if name.startswith("word/media/"):
                info = archive.getinfo(name)
                media.append({"name": name, "size": info.file_size})
        document_xml = archive.read("word/document.xml").decode("utf-8")
        field_summary = {
            "toc": "TOC" in document_xml,
            "page_fields": document_xml.count("PAGE"),
            "section_breaks": document_xml.count("<w:sectPr"),
            "page_breaks": document_xml.count('w:type="page"'),
            "drawings": document_xml.count("<w:drawing"),
        }

    properties = document.core_properties
    return {
        "file": str(path.resolve()),
        "core_properties": {
            "title": properties.title,
            "subject": properties.subject,
            "author": properties.author,
            "last_modified_by": properties.last_modified_by,
            "created": properties.created.isoformat() if properties.created else None,
            "modified": properties.modified.isoformat() if properties.modified else None,
        },
        "counts": {
            "paragraphs": len(document.paragraphs),
            "non_empty_paragraphs": len(paragraphs),
            "tables": len(document.tables),
            "inline_shapes": len(document.inline_shapes),
            "sections": len(document.sections),
            "media": len(media),
        },
        "style_counts": dict(style_counts.most_common()),
        "styles": styles,
        "sections": sections,
        "field_summary": field_summary,
        "media": media,
        "paragraphs": paragraphs,
        "tables": tables,
        "ordered_blocks": ordered_blocks,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect the structure of a DOCX report.")
    parser.add_argument("path", type=Path)
    parser.add_argument("--compact", action="store_true")
    parser.add_argument(
        "--section",
        choices=("all", "summary", "paragraphs", "tables", "blocks", "styles"),
        default="all",
    )
    args = parser.parse_args()
    report = inspect_report(args.path)
    if args.section == "summary":
        payload = {key: value for key, value in report.items() if key not in {"paragraphs", "tables", "ordered_blocks"}}
    elif args.section == "all":
        payload = report
    else:
        key = {"blocks": "ordered_blocks"}.get(args.section, args.section)
        payload = report[key]
    print(json.dumps(payload, ensure_ascii=False, indent=None if args.compact else 2))


if __name__ == "__main__":
    main()
