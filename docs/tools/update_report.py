from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
REPORT = DOCS / "Lưu Phạm Vĩnh Tùng - Đồ án cơ sở - DẤU VỊ.docx"
ORIGINAL = DOCS / "Lưu Phạm Vĩnh Tùng - Đồ án cơ sở - DẤU VỊ - backup trước cập nhật 2026-07-19.docx"
DIAGRAMS = DOCS / "report-assets" / "diagrams"
SCREENSHOTS = DOCS / "report-assets" / "screenshots"


def set_font(run, size: float = 13, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def find_paragraph(document: DocumentObject, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Không tìm thấy đoạn: {text}")


def clear_between(start: Paragraph, end: Paragraph) -> None:
    current = start._p.getnext()
    while current is not None and current is not end._p:
        following = current.getnext()
        current.getparent().remove(current)
        current = following


def add_paragraph(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    paragraph = anchor.insert_paragraph_before(style=style)
    run = paragraph.add_run(text)
    set_font(run)
    if style == "Normal":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_heading(anchor: Paragraph, text: str, level: int) -> Paragraph:
    style = "CAP2" if level == 2 else "CAP3"
    paragraph = add_paragraph(anchor, text, style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = add_paragraph(anchor, f"• {text}")
    paragraph.paragraph_format.left_indent = Cm(1)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    return paragraph


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def style_table(table) -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        row_properties.append(no_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_properties.append(repeat)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade_cell(cell, "E4EEE7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, 13, bold=row_index == 0)


def add_table(
    document: DocumentObject,
    anchor: Paragraph,
    caption: str,
    rows: list[list[str]],
) -> None:
    caption_paragraph = add_paragraph(anchor, caption, "VN_TableCaption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            table.cell(row_index, column_index).text = value
    style_table(table)
    anchor._p.addprevious(table._tbl)
    anchor.insert_paragraph_before()


def add_figure(anchor: Paragraph, image: Path, caption: str) -> None:
    figure = anchor.insert_paragraph_before()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.keep_with_next = True
    figure.add_run().add_picture(str(image), width=Cm(15))
    caption_paragraph = add_paragraph(anchor, caption, "VN_Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code(anchor: Paragraph, text: str) -> None:
    paragraph = anchor.insert_paragraph_before(style="No Spacing")
    run = paragraph.add_run(text)
    set_font(run, 13)
    paragraph.paragraph_format.left_indent = Cm(0.8)


def write_chapter_four(document: DocumentObject) -> None:
    chapter = find_paragraph(document, "CHƯƠNG 4. CÀI ĐẶT CHƯƠNG TRÌNH VÀ TRIỂN KHAI")
    conclusion = find_paragraph(document, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    clear_between(chapter, conclusion)

    add_paragraph(
        conclusion,
        "Chương này trình bày kết quả cài đặt và phương án triển khai website DẤU VỊ. Hệ thống được tổ chức theo kiến trúc tách frontend khỏi backend và cơ sở dữ liệu: frontend Next.js triển khai trên Vercel; FastAPI và PostgreSQL chạy bằng Docker Compose trên máy chủ cloud; DuckDNS và Caddy do người vận hành cấu hình riêng để cung cấp tên miền và HTTPS cho API.",
    )

    add_heading(conclusion, "4.1. Mô tả hệ thống đã cài đặt", 2)
    add_heading(conclusion, "4.1.1. Frontend Next.js", 3)
    add_paragraph(
        conclusion,
        "Frontend sử dụng Next.js, React và TypeScript, được tổ chức theo từng feature. Các trang giao diện làm việc qua repository interfaces thay vì đọc trực tiếp dữ liệu mock. Biến NEXT_PUBLIC_DATA_SOURCE cho phép chuyển giữa mock repositories và HTTP repositories mà không thay đổi component hoặc page.",
    )
    add_bullet(conclusion, "Trang chủ giới thiệu bộ sưu tập và lô cà phê nổi bật.")
    add_bullet(conclusion, "Trang bộ sưu tập hỗ trợ tìm kiếm, lọc và sắp xếp sáu sản phẩm.")
    add_bullet(conclusion, "Trang chi tiết cho phép chọn variant, quy cách xay và thêm vào giỏ hàng.")
    add_bullet(conclusion, "Trang truy xuất hiển thị passport, timeline sáu bước và mức bằng chứng của từng thuộc tính.")
    add_bullet(conclusion, "Coffee Advisor thu thập sở thích qua sáu bước và hiển thị tối đa ba gợi ý có giải thích.")
    add_bullet(conclusion, "Giỏ hàng lưu trên localStorage; cart chỉ được xóa sau khi backend xác nhận tạo đơn thành công.")

    add_heading(conclusion, "4.1.2. Backend FastAPI", 3)
    add_paragraph(
        conclusion,
        "Backend được xây dựng bằng FastAPI, Pydantic và SQLAlchemy. Hệ thống cung cấp REST API phiên bản /api/v1 cho catalog, hồ sơ lô, Coffee Advisor và tạo đơn COD trình diễn. Middleware giới hạn CORS và host, nén response, gắn request ID và security headers. Endpoint readiness thực hiện truy vấn PostgreSQL để phản ánh đúng trạng thái sẵn sàng của toàn dịch vụ.",
    )
    add_table(
        document,
        conclusion,
        "Bảng 4.1. Danh sách endpoint chính của backend",
        [
            ["Phương thức", "Endpoint", "Chức năng"],
            ["GET", "/health/live", "Kiểm tra tiến trình backend đang hoạt động"],
            ["GET", "/health/ready", "Kiểm tra backend và kết nối PostgreSQL"],
            ["GET", "/api/v1/products", "Tìm kiếm, lọc và sắp xếp sản phẩm"],
            ["GET", "/api/v1/products/featured", "Lấy danh sách sản phẩm nổi bật"],
            ["GET", "/api/v1/products/{slug}", "Lấy chi tiết sản phẩm và variant"],
            ["GET", "/api/v1/lots/featured", "Lấy các lô truy xuất nổi bật"],
            ["GET", "/api/v1/lots/{lotCode}", "Lấy passport, timeline và evidence của lô"],
            ["POST", "/api/v1/advisor/recommendations", "Tạo tối đa ba gợi ý có điểm và lý do"],
            ["POST", "/api/v1/orders", "Tạo đơn COD có kiểm tra và chống lặp"],
        ],
    )

    add_heading(conclusion, "4.1.3. PostgreSQL, migration và seed", 3)
    add_paragraph(
        conclusion,
        "PostgreSQL lưu ba nhóm dữ liệu: catalog gồm products và product_variants; truy xuất gồm coffee_lots, evidence_items và lot_timeline_events; đặt hàng gồm orders và order_items. Alembic quản lý phiên bản schema. Khi container backend khởi động, entrypoint chạy migration tới head rồi thực hiện seed idempotent sáu sản phẩm và sáu lô demo.",
    )
    add_paragraph(
        conclusion,
        "Order item giữ snapshot tên sản phẩm, SKU, định dạng và đơn giá tại thời điểm đặt hàng. Cách thiết kế này giúp đơn cũ vẫn đọc đúng ngay cả khi catalog thay đổi sau đó. Database chỉ nằm trên mạng Docker nội bộ và không publish cổng ra máy chủ.",
    )

    add_heading(conclusion, "4.1.4. Coffee Advisor", 3)
    add_paragraph(
        conclusion,
        "Coffee Advisor sử dụng bộ luật xác định trên catalog hiện có. Pydantic kiểm tra preferences; service loại các variant sai định dạng, hết hàng hoặc vượt ngân sách; sau đó tính điểm theo cách pha, body, độ đắng, độ chua, caffeine, giá và priorities. Điểm được chuẩn hóa về thang 0–100 và mỗi kết quả có tối đa bốn lý do.",
    )

    add_heading(conclusion, "4.1.5. Checkout COD", 3)
    add_paragraph(
        conclusion,
        "Frontend chỉ gửi productId, variantId, grind và quantity, không gửi unitPrice. Backend đọc lại variant hiện tại, kiểm tra tồn kho và grind option, tính line total, subtotal và phí giao hàng. Phí giao hàng là 30.000 đồng khi subtotal dưới 499.000 đồng. Orders và order_items được ghi trong một transaction; Idempotency-Key giúp request lặp không tạo thêm đơn trùng.",
    )

    add_heading(conclusion, "4.2. Giao diện chương trình", 2)
    add_paragraph(
        conclusion,
        "Các hình sau được chụp trực tiếp từ production build của frontend ở độ rộng 1440 px. Dữ liệu truy xuất trong giao diện được gắn nhãn demo, tránh gây hiểu nhầm đây là dữ liệu chứng nhận thực tế.",
    )
    add_figure(conclusion, SCREENSHOTS / "01-home.png", "Hình 4.1. Giao diện trang chủ DẤU VỊ")
    add_figure(conclusion, SCREENSHOTS / "02-shop.png", "Hình 4.2. Giao diện bộ sưu tập và bộ lọc")
    add_figure(conclusion, SCREENSHOTS / "03-traceability.png", "Hình 4.3. Giao diện passport truy xuất theo mã lô")
    add_figure(conclusion, SCREENSHOTS / "04-advisor.png", "Hình 4.4. Giao diện Coffee Advisor")

    add_heading(conclusion, "4.3. Đóng gói backend và database bằng Docker", 2)
    add_paragraph(
        conclusion,
        "Tệp docker/compose.yml định nghĩa hai service backend và database. PostgreSQL sử dụng named volume để giữ dữ liệu và healthcheck bằng pg_isready. Backend chạy bằng user không có quyền root, filesystem chỉ đọc, tmpfs cho thư mục tạm và tùy chọn no-new-privileges. Cổng backend chỉ bind 127.0.0.1:8000 để Caddy trên cùng máy chủ reverse proxy vào.",
    )
    add_figure(conclusion, DIAGRAMS / "05-cloud-deployment.png", "Hình 4.5. Quy trình triển khai độc lập frontend và backend/database")
    add_paragraph(
        conclusion,
        "Trước khi khởi động, sao chép docker/.env.example thành docker/.env và thay toàn bộ giá trị mẫu, đặc biệt là mật khẩu PostgreSQL, domain Vercel và hostname API. Các lệnh chính được thực hiện tại thư mục gốc của repository:",
    )
    add_code(conclusion, "Copy-Item docker/.env.example docker/.env")
    add_code(conclusion, "docker compose --env-file docker/.env -f docker/compose.yml config")
    add_code(conclusion, "docker compose --env-file docker/.env -f docker/compose.yml up -d --build")
    add_code(conclusion, "docker compose --env-file docker/.env -f docker/compose.yml ps")
    add_code(conclusion, "Invoke-WebRequest http://127.0.0.1:8000/health/ready")

    add_heading(conclusion, "4.4. Triển khai Vercel, DuckDNS và Caddy", 2)
    add_paragraph(
        conclusion,
        "Frontend được triển khai trên Vercel với Root Directory là frontend. Sau khi backend hoạt động qua HTTPS, cấu hình các biến NEXT_PUBLIC_API_BASE_URL, NEXT_PUBLIC_DATA_SOURCE=http và NEXT_PUBLIC_ENABLE_CHECKOUT=true rồi redeploy frontend. NEXT_PUBLIC_API_BASE_URL phải kết thúc tại /api/v1.",
    )
    add_paragraph(
        conclusion,
        "DuckDNS và Caddy không nằm trong Compose vì đã được người vận hành cài đặt riêng. Caddy nhận request HTTPS từ hostname API và reverse_proxy tới 127.0.0.1:8000. CORS_ORIGINS phía backend phải khớp chính xác domain Vercel; ALLOWED_HOSTS phải chứa hostname DuckDNS dùng cho API.",
    )
    add_table(
        document,
        conclusion,
        "Bảng 4.2. Các biến môi trường quan trọng khi triển khai",
        [
            ["Nơi cấu hình", "Biến", "Giá trị mẫu hoặc ý nghĩa"],
            ["Vercel", "NEXT_PUBLIC_API_BASE_URL", "https://api.<duckdns-domain>/api/v1"],
            ["Vercel", "NEXT_PUBLIC_DATA_SOURCE", "http"],
            ["Vercel", "NEXT_PUBLIC_ENABLE_CHECKOUT", "true"],
            ["Backend", "CORS_ORIGINS", "Domain frontend Vercel được phép gọi API"],
            ["Backend", "ALLOWED_HOSTS", "Hostname API, localhost và 127.0.0.1"],
            ["PostgreSQL", "POSTGRES_PASSWORD", "Mật khẩu production ngẫu nhiên và đủ dài"],
            ["Compose", "BACKEND_BIND_ADDRESS", "127.0.0.1 khi Caddy chạy cùng máy chủ"],
        ],
    )

    add_heading(conclusion, "4.5. Kiểm thử và kết quả", 2)
    add_paragraph(
        conclusion,
        "Hệ thống được kiểm tra ở mức unit/API, kiểm tra kiểu dữ liệu, production build và chạy tích hợp bằng container thật. Luồng Docker được xác minh với PostgreSQL mới hoàn toàn: migration tạo schema, seed tạo dữ liệu, readiness trả trạng thái sẵn sàng, API trả đủ catalog và backend khởi động lại không tạo dữ liệu trùng.",
    )
    add_table(
        document,
        conclusion,
        "Bảng 4.3. Kết quả kiểm thử tại thời điểm hoàn thiện",
        [
            ["Hạng mục", "Kết quả", "Phạm vi kiểm tra"],
            ["Backend pytest", "6/6 test đạt", "Catalog, lot, Advisor, order và idempotency"],
            ["Backend Ruff", "Đạt", "Lint và format mã Python"],
            ["Frontend Vitest", "24/24 test, 8 file đạt", "Domain, repository, filter, Advisor và checkout API"],
            ["Frontend ESLint", "Đạt", "Quy tắc chất lượng mã nguồn"],
            ["TypeScript typecheck", "Đạt", "Kiểm tra kiểu dữ liệu toàn frontend"],
            ["Next.js production build", "Đạt, 18 route", "Static, SSG và dynamic route"],
            ["Docker Compose", "Đạt", "Build image, migration, seed, health và PostgreSQL thật"],
        ],
    )

    add_heading(conclusion, "4.6. Hạn chế của phiên bản hiện tại", 2)
    add_bullet(conclusion, "Chưa có tài khoản, phân quyền, trang quản trị và lịch sử đơn hàng cho khách.")
    add_bullet(conclusion, "Checkout là COD trình diễn; chưa tích hợp thanh toán, email, SMS hoặc đơn vị vận chuyển.")
    add_bullet(conclusion, "Tồn kho được kiểm tra nhưng chưa trừ kho và chưa có cơ chế giữ hàng đồng thời.")
    add_bullet(conclusion, "Dữ liệu truy xuất là dữ liệu demo, không thay thế chứng nhận từ nhà sản xuất.")
    add_bullet(conclusion, "Coffee Advisor sử dụng luật xác định, chưa tích hợp LLM hoặc RAG.")


def update_report(source: Path, output: Path) -> None:
    if not source.exists():
        raise FileNotFoundError(f"Không tìm thấy bản gốc: {source}")
    assets = [
        DIAGRAMS / "05-cloud-deployment.png",
        SCREENSHOTS / "01-home.png",
        SCREENSHOTS / "02-shop.png",
        SCREENSHOTS / "03-traceability.png",
        SCREENSHOTS / "04-advisor.png",
    ]
    missing = [str(path) for path in assets if not path.exists()]
    if missing:
        raise FileNotFoundError("Thiếu tài nguyên báo cáo:\n" + "\n".join(missing))

    document = Document(source)
    write_chapter_four(document)
    temporary = output.with_name(f".{output.stem}.chapter4.docx")
    document.save(temporary)
    Document(temporary)
    temporary.replace(output)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Khôi phục bản gốc và chỉ cập nhật nội dung bên trong Chương 4."
    )
    parser.add_argument("--source", type=Path, default=ORIGINAL)
    parser.add_argument("--output", type=Path, default=REPORT)
    args = parser.parse_args()
    update_report(args.source.resolve(), args.output.resolve())
    print(f"Đã cập nhật riêng Chương 4: {args.output.resolve()}")
    print(f"Nguồn giữ nguyên: {args.source.resolve()}")


if __name__ == "__main__":
    main()
